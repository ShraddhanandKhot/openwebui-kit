import os
import uuid
import json
import re
import sqlite3
import time
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from open_webui.models.files import Files, FileForm
except ImportError:
    from open_webui.apps.webui.models.files import Files, FileForm


class Tools:

    def __init__(self):
        pass

    # ============================================================
    # INSTANT KB LINKER
    # ============================================================

    def _link_to_output_kb(self, file_id: str, user_id: str) -> None:
        """
        Instantly links the generated file to the 'output' Knowledge Base in SQLite.
        """
        db_path = "./data/webui.db"
        try:
            con = sqlite3.connect(db_path)
            cur = con.cursor()

            # 1. Fetch or create 'output' KB ID
            cur.execute("SELECT id FROM knowledge WHERE name='output'")
            row = cur.fetchone()

            if row:
                kb_id = row[0]
            else:
                kb_id = str(uuid.uuid4())
                now = int(time.time())
                cur.execute(
                    "INSERT INTO knowledge (id, user_id, name, description, created_at, updated_at) "
                    "VALUES (?, ?, 'output', 'Auto-collected OpenWebUI-generated files', ?, ?)",
                    (kb_id, user_id, now, now),
                )

            # 2. Insert link row into knowledge_file table instantly
            now = int(time.time())
            cur.execute(
                "INSERT OR IGNORE INTO knowledge_file "
                "(id, user_id, knowledge_id, file_id, created_at, updated_at, directory_id) "
                "VALUES (?, ?, ?, ?, ?, ?, NULL)",
                (str(uuid.uuid4()), user_id, kb_id, file_id, now, now),
            )
            con.commit()
            con.close()
        except Exception as e:
            print(f"Failed to instantly link file to output KB: {e}")

    # ============================================================
    # PAGE BORDER
    # ============================================================

    def _add_page_border(self, section):
        sectPr = section._sectPr
        existing = sectPr.find(qn("w:pgBorders"))
        if existing is not None:
            sectPr.remove(existing)

        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")

        for side in ["top", "bottom", "left", "right"]:
            elem = OxmlElement(f"w:{side}")
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "8")
            elem.set(qn("w:space"), "18")
            elem.set(qn("w:color"), "808080")
            pgBorders.append(elem)

        sectPr.append(pgBorders)

    # ============================================================
    # FOOTER
    # ============================================================

    def _create_footer(self, section):
        footer = section.footer
        for paragraph in footer.paragraphs:
            paragraph.clear()

        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run("OpenWebUI")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)

        run = paragraph.add_run("\n")

        run = paragraph.add_run("Spryntworks")
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # ============================================================
    # DOCUMENT CONFIGURATION
    # ============================================================

    def _configure_document(self, document):
        for section in document.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)
            section.footer_distance = Inches(0.35)
            self._create_footer(section)
            self._add_page_border(section)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
            try:
                style = document.styles[style_name]
                style.font.name = "Arial"
                style.font.size = Pt(size)
                style.font.bold = True
            except Exception:
                pass

    def _add_title(self, document, title):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(18)
        run = paragraph.add_run(title)
        run.bold = True
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(31, 78, 121)

    def _add_heading(self, document, text, level=1):
        paragraph = document.add_heading(text, level=level)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    def _add_paragraph(self, document, text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(str(text))
        return paragraph

    def _add_bullets(self, document, items):
        for item in items:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(str(item))

    def _add_numbered_list(self, document, items):
        for item in items:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(str(item))

    def _add_table(self, document, table_data):
        columns = table_data.get("columns", [])
        rows = table_data.get("rows", [])
        if not columns:
            return

        table = document.add_table(rows=1, cols=len(columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for index, column in enumerate(columns):
            cell = header_cells[index]
            cell.text = str(column)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

        for row_data in rows:
            row_cells = table.add_row().cells
            for index in range(len(columns)):
                value = row_data[index] if index < len(row_data) else ""
                row_cells[index].text = str(value)
                row_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in row_cells[index].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9)

        document.add_paragraph()

    def _add_content(self, document, content):
        sections = content.get("sections", [])
        for section_data in sections:
            if not isinstance(section_data, dict):
                continue

            heading = section_data.get("heading")
            level = section_data.get("level", 1)
            if heading:
                self._add_heading(document, heading, level)

            paragraphs = section_data.get("paragraphs", [])
            for paragraph in paragraphs:
                self._add_paragraph(document, paragraph)

            bullets = section_data.get("bullets", [])
            if bullets:
                self._add_bullets(document, bullets)

            numbered = section_data.get("numbered", [])
            if numbered:
                self._add_numbered_list(document, numbered)

            tables = section_data.get("tables", [])
            for table_data in tables:
                if isinstance(table_data, dict):
                    self._add_table(document, table_data)

    def _chat_to_content(self, messages):
        sections = []
        if not messages:
            return sections

        for index, message in enumerate(messages, start=1):
            if not isinstance(message, dict):
                continue

            role = message.get("role", "unknown")
            content = message.get("content", "")

            if not isinstance(content, str):
                try:
                    content = json.dumps(content, ensure_ascii=False)
                except Exception:
                    content = str(content)

            content = content.strip()
            if not content:
                continue

            role_title = role.capitalize()
            sections.append(
                {
                    "heading": f"{role_title} Message {index}",
                    "level": 2,
                    "paragraphs": [content],
                }
            )

        return sections

    # ============================================================
    # MAIN TOOL METHOD
    # ============================================================

    async def create_docx_file(
        self,
        title: str,
        content: str = "",
        filename: str = "generated_document.docx",
        __user__: dict = None,
        __messages__: list = None,
    ) -> dict:
        """
        Generate a formatted Word (.docx) document with custom styling.

        :param title: Main title displayed at the top of the document.
        :param content: Structured JSON string containing sections, headings, paragraphs, bullet points, and tables. If exporting chat, pass 'this chat'.
        :param filename: Desired filename ending in .docx (e.g., 'report.docx').
        :return: Structured result dictionary containing file attachment metadata.
        """
        content_lower = content.lower() if content else ""
        chat_request = any(
            phrase in content_lower
            for phrase in [
                "this chat",
                "current chat",
                "this conversation",
                "current conversation",
                "convert chat",
                "convert this chat",
                "export chat",
                "export this chat",
            ]
        )

        document = Document()
        self._configure_document(document)
        self._add_title(document, title)

        if chat_request and __messages__:
            sections = self._chat_to_content(__messages__)
            structured_content = {"sections": sections}
            self._add_content(document, structured_content)
        else:
            if not content:
                return {"error": "No document content was provided."}

            try:
                structured_content = json.loads(content)
            except Exception:
                paragraphs = [
                    line.strip() for line in content.splitlines() if line.strip()
                ]
                structured_content = {
                    "sections": [
                        {"heading": "Content", "level": 1, "paragraphs": paragraphs}
                    ]
                }

            self._add_content(document, structured_content)

        document.add_paragraph()
        info_paragraph = document.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = info_paragraph.add_run("Generated by Open WebUI")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(120, 120, 120)

        if not filename.lower().endswith(".docx"):
            filename = filename.rsplit(".", 1)[0] + ".docx"

        filename = re.sub(r'[<>:"/\\|?*]', "_", filename)

        upload_dir = os.path.abspath("./data/uploads")
        os.makedirs(upload_dir, exist_ok=True)

        file_id = str(uuid.uuid4())
        unique_filename = f"{file_id}_{filename}"
        file_path = os.path.join(upload_dir, unique_filename)

        document.save(file_path)

        user_id = __user__.get("id") if __user__ else "admin"
        file_size = os.path.getsize(file_path)

        # Register file with Open WebUI database
        file_form = FileForm(
            id=file_id,
            filename=filename,
            path=file_path,
            data={"title": title, "content": content},
            meta={
                "name": filename,
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "size": file_size,
                "description": "DOCX document generated by Open WebUI",
                "tags": ["docx", "word", "document", "generated"],
            },
        )

        result = await Files.insert_new_file(user_id, file_form)

        if result is None:
            return {"error": "DOCX was created but failed to register inside Open WebUI."}

        # INSTANTLY LINK TO 'output' KNOWLEDGE BASE
        self._link_to_output_kb(file_id, user_id)

        # Return structured output format for chat auto-attachment
        return {
            "output": f"DOCX document created successfully. Filename: {filename}",
            "files": [
                {
                    "id": file_id,
                    "filename": filename,
                    "name": filename,
                    "url": file_id,
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "size": file_size,
                    "type": "file",
                    "status": "processed",
                }
            ],
        }
