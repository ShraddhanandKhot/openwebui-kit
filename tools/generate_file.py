import os
import re
import uuid

from io import BytesIO

try:
    from fastapi import UploadFile, Headers
except ImportError:
    from fastapi import UploadFile
    from starlette.datastructures import Headers


def _safe_name(title: str) -> str:
    name = re.sub(r"[^\w\- ]+", "", title or "file").strip().replace(" ", "_")
    return (name or "file")[:80]


def _build_pdf(title: str, content: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, title[:120])
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    # core fonts are latin-1 only; replace anything else so it never crashes
    text = content.encode("latin-1", errors="replace").decode("latin-1")
    for para in text.split("\n\n"):
        pdf.multi_cell(0, 6, para.strip())
        pdf.ln(2)
    return pdf.output()


def _build_docx(title: str, content: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title[:120], level=0)
    for para in content.split("\n\n"):
        doc.add_paragraph(para.strip().replace("\n", " "))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx(title: str, content: str) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        if "\t" in line:
            cells = [c.strip() for c in line.split("\t")]
        else:
            cells = [c.strip() for c in line.split(",")]
        ws.append(cells[:50])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


_BUILDERS = {"pdf": _build_pdf, "docx": _build_docx, "xlsx": _build_xlsx}
_MIMES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


class Tools:
    def __init__(self):
        pass

    async def generate_file(
        self,
        format: str,
        title: str,
        content: str,
        __user__: dict = None,
        __event_emitter__=None,
        __request__=None,
    ) -> str:
        """Generate a downloadable file (PDF, DOCX, or XLSX) from the given content, register it in Open WebUI Files, and post its download link in the chat. Use this tool whenever the user asks to generate, create, export, or download a file such as a report, document, or spreadsheet. Parameters: format is one of pdf, docx, xlsx; title is the filename without extension (e.g. Weekly_Summary); content is the full text/data for the file (for xlsx use one row per line with tab-separated or comma-separated columns; for pdf/docx use plain text or markdown). Do NOT fabricate a file yourself - always call this tool."""

        fmt = (format or "").strip().lower().lstrip(".")
        if fmt not in _BUILDERS:
            return "Invalid format. Use one of: pdf, docx, xlsx. Tell the user the file could not be generated."

        if not title or not title.strip():
            title = "generated_file"
        filename = f"{_safe_name(title)}.{fmt}"

        try:
            data = _BUILDERS[fmt](title, content)
        except Exception as exc:
            return f"Failed to build the file: {exc}"

        url = None
        file_id = None

        # ---- register through Open WebUI's real upload path so /content serves it ----
        try:
            from open_webui.routers.files import upload_file_handler
            from open_webui.models.users import Users

            user_model = await Users.get_user_by_id(__user__["id"]) if __user__ else None
            if user_model is not None and __request__ is not None:
                upload = UploadFile(
                    file=BytesIO(data),
                    filename=filename,
                    headers=Headers({"content-type": _MIMES[fmt]}),
                )
                item = await upload_file_handler(
                    request=__request__,
                    file=upload,
                    metadata={},
                    process=False,
                    user=user_model,
                )
                if item and getattr(item, "id", None):
                    file_id = item.id
                    url = f"/api/v1/files/{file_id}/content"
        except Exception:
            url = None

        # ---- fallback: static cache dir (download works, no Files-panel entry) ----
        if url is None:
            try:
                stored = f"{_safe_name(title)}_{uuid.uuid4().hex[:6]}.{fmt}"
                path = os.path.join("/app/backend/data/cache/files", stored)
                os.makedirs(os.path.dirname(path), exist_ok=True)
                with open(path, "wb") as fh:
                    fh.write(data)
                url = f"/cache/files/{stored}"
            except Exception:
                return f"Generated {filename} but could not save it for download."

        if file_id:
            # Return the file object; the patched middleware (process_tool_result)
            # attaches it to the chat message natively as a downloadable file.
            return {
                "output": f"Generated {filename}",
                "files": [
                    {
                        "id": file_id,
                        "filename": filename,
                        "name": filename,
                        "url": file_id,
                        "content_type": _MIMES[fmt],
                        "size": len(data),
                        "type": "file",
                        "status": "processed",
                    }
                ],
            }
        return f"Generated {filename} but could not attach it to the chat."
